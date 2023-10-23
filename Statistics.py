import os
import sys
from tkinter import *
import docx
import matplotlib.pyplot as plt
import numpy as np
from PIL import Image, ImageTk
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import geopandas as gpd


class Aggregation:

    def __init__(self, db_manager):
        self.db_manager = db_manager

    def draw_chart(self, dictionary, country, plot_type):
        filtered_data = {key: value for key, value in dictionary.items() if value != 0}
        print("In draw function")
        x = list(filtered_data.keys())
        y = list(filtered_data.values())

        plt.ylim(0, max(y))
        plt.figure(figsize=(10, 6))
        plt.xlabel("Year")
        plt.ylabel("Population")
        plt.title("Population by year in " + country)
        plt.ticklabel_format(style='plain')  # Disable scaling on y-axis

        if plot_type == "Bar chart":
            plt.bar(x, y)
        elif plot_type == "Line chart":
            plt.plot(x, y)
        elif plot_type == "Pie chart":
            plt.pie(y, labels=x)
        else:
            raise ValueError("Invalid chart type provided.")

        plt.savefig("plot.png")

    def draw_general_plot(self, dict, plot_type, year):
        filtered_data = {key: value for key, value in dict.items() if value != 0}
        x = list(filtered_data.keys())
        y = [float(value) for value in filtered_data.values()]

        plt.figure(figsize=(10, 6))
        plt.ylim(0, max(y))
        plt.ylabel("Population")
        plt.xticks([])  # Remove x-axis labels
        plt.yticks(np.arange(10_000_000, max(y) + max(y) / 10, 50_000_000))
        plt.ticklabel_format(style='plain')  # Disable scaling on y-axis

        plt.title(f"Population by year in {year}")

        if plot_type == "Bar chart":
            plt.bar(x, y)
        elif plot_type == "Line chart":
            plt.plot(x, y)
        elif plot_type == "Pie chart":
            plt.pie(y, labels=x)
        else:
            raise ValueError("Invalid chart type provided.")

        plt.savefig("plot.png")

    def create_geoplot(self, dictionary):
        shapefile_path = r'E:\PyCharm Community Edition 2022.3.2\lab12_upd\geo\ne_50m_admin_0_countries.shp'
        gdf = gpd.read_file(shapefile_path)

        # Step 3: Join the population data with the GeoDataFrame
        gdf['population'] = gdf['ADMIN'].map(dictionary)
        # Step 4: Visualize the map using GeoPandas and matplotlib
        gdf.plot(column='population', cmap='Blues', linewidth=0.8, edgecolor='0.8', figsize=(12, 8))
        plt.title('Population Map')
        plt.savefig("plot.png")

    def resize_image(self, imagename, width, height):
        image = Image.open(imagename)
        resized_image = image.resize((width, height))

        tk_image = ImageTk.PhotoImage(resized_image)

        return tk_image

    def calculate_aggregation(self, arguments):
        cursor = self.db_manager.get_connection_cursor()

        def average_population():
            years = [1980, 2000, 2010, 2022, 2023, 2030, 2050]
            avg_population_dict = {year: 0 for year in years}
            total = 0

            for year in years:
                cursor.execute(f"SELECT pop{year} FROM my_table")
                rows = cursor.fetchall()
                i = 0
                for row in rows:
                    total += int(row[0])
                    i += 1

                avg_population_dict[year] = int(total / i)
            return avg_population_dict

        def maximal_population():
            years = [1980, 2000, 2010, 2022, 2023, 2030, 2050]
            population_dict = {year: 0 for year in years}

            for year in years:
                max_val = 0
                country = ""
                cursor.execute(f"SELECT pop{year}, country FROM my_table")
                rows = cursor.fetchall()
                for row in rows:
                    if int(row[0]) > max_val:
                        max_val = int(row[0])
                        country = row[1]

                population_dict[year] = f"{max_val} in {country}"
            return population_dict

        def minimal_population():
            years = [1980, 2000, 2010, 2022, 2023, 2030, 2050]
            population_dict = {year: 0 for year in years}

            for year in years:
                min_val = sys.maxsize
                min_country = ""
                cursor.execute(f"SELECT pop{year}, country FROM my_table")
                rows = cursor.fetchall()
                for row in rows:
                    if int(row[0]) < min_val:
                        min_val = int(row[0])
                        min_country = row[1]

                population_dict[year] = f"{min_val} in {min_country}"
            return population_dict

        def max(arg):
            max_val = 0
            country = ''
            cursor.execute(f"SELECT country, {arg} FROM my_table")
            rows = cursor.fetchall()
            for row in rows:
                if float(row[1]) > max_val:
                    max_val = float(row[1])
                    country = row[0]

            return country, max_val

        def min(arg):
            min_val = sys.maxsize
            country = ''
            cursor.execute(f"SELECT country, {arg} FROM my_table")
            rows = cursor.fetchall()
            for row in rows:
                if float(row[1]) < min_val:
                    min_val = float(row[1])
                    country = row[0]

            return country, min_val

        output = ""
        for arg in arguments:
            if len(arg.split(" ")) == 2:
                function, argument = arg.split(" ")
                if function == "Maximal":
                    output += f"Maximal {argument}: " + str(max(argument)) + "\n"
                elif function == "Minimal":
                    output += f"Minimal {argument}: " + str(min(argument)) + "\n"
            elif arg == "Maximal population by year":
                output += "Maximal population by year: \n"
                for key, value in maximal_population().items():
                    output += f"{key}: {value} \n"
            elif arg == "Minimal population by year":
                output += "Minimal population by year: \n"
                for key, value in minimal_population().items():
                    output += f"{key}: {value} \n"
            elif arg == "Average population by year":
                output += "Average population by year: \n"
                for key, value in average_population().items():
                    output += f"{key}: {value} \n"
        return output


    def generate_report(self, docx_name, paragraph_dict):
        if os.path.exists(docx_name):
            filename, extension = os.path.splitext(docx_name)
            count = 1
            while os.path.exists(f"{filename}({count}){extension}"):
                count += 1

            docx_name = f"{filename}({count}){extension}"

        document = docx.Document()

        title_paragraph = document.add_paragraph()
        title_paragraph_run = title_paragraph.add_run("Countries analyser")
        title_paragraph_run.bold = True
        title_paragraph_run.font.size = docx.shared.Pt(22)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        report_author_paragraph = document.add_paragraph()
        report_author_paragraph_run = report_author_paragraph.add_run("Bohdan Kyryliuk")
        report_author_paragraph_run.font.size = docx.shared.Pt(18)
        report_author_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        author_paragraph = document.add_paragraph()
        author_paragraph_run = author_paragraph.add_run("267855")
        author_paragraph_run.font.size = docx.shared.Pt(18)
        author_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        document.add_page_break()

        plot_paragraph = document.add_paragraph()
        plot_paragraph_run = plot_paragraph.add_run()
        image = Image.open("plot.png")
        resized_image = image.resize((500, 500))
        resized_image.save("resized_plot.png")
        plot_paragraph_run.add_picture("resized_plot.png")

        plot_report = paragraph_dict
        plot_report_paragraph = document.add_paragraph()
        plot_report_paragraph_run = plot_report_paragraph.add_run(plot_report)
        plot_report_paragraph_run.font.size = docx.shared.Pt(18)

        document.save(docx_name)
